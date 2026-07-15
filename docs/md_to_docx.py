# -*- coding: utf-8 -*-
"""将 Markdown 文件转换为 Word 文档"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches


def set_doc_font(doc):
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style.font.size = Pt(10.5)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')


def add_formatted_text(paragraph, text):
    pattern = r'(\*\*[^*]+\*\*|`[^`]+`)'
    parts = re.split(pattern, text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        else:
            paragraph.add_run(part)


def parse_table_row(line):
    line = line.strip()
    if line.startswith('|'):
        line = line[1:]
    if line.endswith('|'):
        line = line[:-1]
    return [cell.strip() for cell in line.split('|')]


def is_table_separator(line):
    return bool(re.match(r'^\|?[\s\-:|]+\|?$', line.strip()))


def convert_md_to_docx(md_path, docx_path):
    content = Path(md_path).read_text(encoding='utf-8')
    lines = content.splitlines()

    doc = Document()
    set_doc_font(doc)

    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.18)
    section.right_margin = Inches(1.18)

    i = 0
    in_code = False
    code_lines = []
    list_buffer = []

    def flush_list():
        nonlocal list_buffer
        for item in list_buffer:
            p = doc.add_paragraph(style='List Bullet')
            add_formatted_text(p, item)
        list_buffer = []

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith('```'):
            if in_code:
                p = doc.add_paragraph()
                run = p.add_run('\n'.join(code_lines))
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
                p.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code = False
            else:
                flush_list()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not stripped:
            flush_list()
            i += 1
            continue

        if stripped == '---':
            flush_list()
            doc.add_paragraph('—' * 40)
            i += 1
            continue

        if stripped.startswith('#'):
            flush_list()
            level = len(stripped) - len(stripped.lstrip('#'))
            text = stripped[level:].strip()
            style = f'Heading {min(level, 3)}'
            p = doc.add_paragraph(style=style)
            add_formatted_text(p, text)
            i += 1
            continue

        if stripped.startswith('|') and i + 1 < len(lines) and is_table_separator(lines[i + 1]):
            flush_list()
            headers = parse_table_row(stripped)
            i += 2
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                rows.append(parse_table_row(lines[i]))
                i += 1
            table = doc.add_table(rows=1 + len(rows), cols=len(headers))
            table.style = 'Table Grid'
            for j, h in enumerate(headers):
                cell = table.rows[0].cells[j]
                cell.text = h
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            for r_idx, row in enumerate(rows):
                for c_idx, val in enumerate(row):
                    if c_idx < len(table.rows[r_idx + 1].cells):
                        table.rows[r_idx + 1].cells[c_idx].text = val
            doc.add_paragraph()
            continue

        if re.match(r'^[-*]\s', stripped):
            item = re.sub(r'^[-*]\s+', '', stripped)
            list_buffer.append(item)
            i += 1
            continue

        if re.match(r'^\d+\.\s', stripped):
            flush_list()
            item = re.sub(r'^\d+\.\s+', '', stripped)
            p = doc.add_paragraph(style='List Number')
            add_formatted_text(p, item)
            i += 1
            continue

        if stripped.startswith('> '):
            flush_list()
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            add_formatted_text(p, stripped[2:])
            i += 1
            continue

        flush_list()
        p = doc.add_paragraph()
        add_formatted_text(p, stripped)
        i += 1

    flush_list()
    doc.save(docx_path)
    print(f'已生成: {docx_path}')


if __name__ == '__main__':
    base = Path(__file__).parent
    files = [
        ('PRD.md', '台州直饮水运维系统-PRD.docx'),
        ('产品功能说明书.md', '台州直饮水运维系统-产品功能说明书.docx'),
    ]
    for md_name, docx_name in files:
        convert_md_to_docx(base / md_name, base / docx_name)
